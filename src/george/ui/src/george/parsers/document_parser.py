"""Multi-format document parser for .docx, .md, and .txt files."""
import os
import codecs
import chardet
from typing import Optional, Dict, Any

# Try to import optional dependencies
try:
    import magic  # python-magic for file type detection
    MAGIC_AVAILABLE = True
except (ImportError, OSError):
    MAGIC_AVAILABLE = False
    magic = None

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown = None
class DocumentParserError(Exception):
    """Custom exception for document parsing errors."""
    pass
class DocumentParser:
    """A robust multi-format document parser for .docx, .md, and .txt files."""
    def __init__(self):
        """Initialize the document parser."""
        self.supported_formats = {
            '.txt': 'Plain Text',
            '.md': 'Markdown',
            '.docx': 'Microsoft Word Document'
        }
    def detect_file_type(self, file_path: str) -> str:
        """
        Detect the file type using python-magic or file extension.
        Args:
            file_path (str): Path to the file
        Returns:
            str: Detected file type (extension)
        """
        if not os.path.exists(file_path):
            raise DocumentParserError(f"File not found: {file_path}")
        # Try to detect using python-magic first
        if MAGIC_AVAILABLE:
            try:
                mime = magic.from_file(file_path, mime=True)
                if mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    return '.docx'
                elif mime == 'text/plain':
                    return '.txt'
                # For markdown, we'll need to check extension as it's text/plain
            except Exception:
                pass
        # Fallback to file extension
        _, ext = os.path.splitext(file_path.lower())
        return ext
    def detect_encoding(self, file_path: str) -> str:
        """
        Detect the text encoding of a file.
        Args:
            file_path (str): Path to the file
        Returns:
            str: Detected encoding
        """
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                return result['encoding'] or 'utf-8'
        except Exception as e:
            raise DocumentParserError(f"Error detecting encoding: {str(e)}")
    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a .docx file and extract text content.
        Args:
            file_path (str): Path to the .docx file
        Returns:
            Dict[str, Any]: Parsed content with metadata
        """
        if not DOCX_AVAILABLE:
            raise DocumentParserError("python-docx library not available. Install with: pip install python-docx")
        try:
            doc = DocxDocument(file_path)
            paragraphs = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    paragraphs.append(paragraph.text)
            # Extract basic metadata
            core_properties = doc.core_properties
            metadata = {
                'title': core_properties.title,
                'author': core_properties.author,
                'created': core_properties.created,
                'modified': core_properties.modified,
                'paragraph_count': len(doc.paragraphs)
            }
            return {
                'content': '\n'.join(paragraphs),
                'metadata': metadata,
                'format': 'docx'
            }
        except Exception as e:
            raise DocumentParserError(f"Error parsing .docx file: {str(e)}")
    def parse_markdown(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a .md file and extract text content.
        Args:
            file_path (str): Path to the .md file
        Returns:
            Dict[str, Any]: Parsed content with metadata
        """
        if not MARKDOWN_AVAILABLE:
            raise DocumentParserError("markdown library not available. Install with: pip install markdown")
        try:
            encoding = self.detect_encoding(file_path)
            with codecs.open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            # Convert markdown to plain text (simple conversion)
            # In a real implementation, you might want to use a proper markdown parser
            # that can extract just the text content
            lines = content.split('\n')
            text_lines = []
            for line in lines:
                # Remove markdown formatting characters but keep content
                clean_line = line
                # Remove headers
                if clean_line.startswith('#'):
                    clean_line = clean_line.lstrip('# ')
                # Remove emphasis markers
                clean_line = clean_line.replace('*', '').replace('_', '')
                # Remove inline code markers
                clean_line = clean_line.replace('`', '')
                if clean_line.strip():
                    text_lines.append(clean_line)
            return {
                'content': '\n'.join(text_lines),
                'metadata': {
                    'line_count': len(lines),
                    'encoding': encoding
                },
                'format': 'markdown'
            }
        except Exception as e:
            raise DocumentParserError(f"Error parsing .md file: {str(e)}")
    def parse_txt(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a .txt file and extract text content.
        Args:
            file_path (str): Path to the .txt file
        Returns:
            Dict[str, Any]: Parsed content with metadata
        """
        try:
            encoding = self.detect_encoding(file_path)
            with codecs.open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            lines = content.split('\n')
            return {
                'content': content,
                'metadata': {
                    'line_count': len(lines),
                    'encoding': encoding
                },
                'format': 'txt'
            }
        except Exception as e:
            raise DocumentParserError(f"Error parsing .txt file: {str(e)}")
    def parse(self, file_path: str) -> Dict[str, Any]:
        """
        Parse a document file based on its format.
        Args:
            file_path (str): Path to the document file
        Returns:
            Dict[str, Any]: Parsed content with metadata
        """
        if not os.path.exists(file_path):
            raise DocumentParserError(f"File not found: {file_path}")
        # Detect file type
        file_type = self.detect_file_type(file_path)
        # Parse based on file type
        if file_type == '.docx':
            return self.parse_docx(file_path)
        elif file_type == '.md':
            return self.parse_markdown(file_path)
        elif file_type == '.txt':
            return self.parse_txt(file_path)
        else:
            # Try to parse as text if extension is unknown
            try:
                return self.parse_txt(file_path)
            except DocumentParserError:
                raise DocumentParserError(f"Unsupported file format: {file_type}")