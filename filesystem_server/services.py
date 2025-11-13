"""Core services for the filesystem server."""
import os
import codecs
import chardet
import re
import logging
import requests
from typing import List, Dict, Any, Optional, Tuple
from dataclasses import dataclass
from bs4 import BeautifulSoup  # For WebSanitizer

# Optional dependencies
try:
    import magic
    MAGIC_AVAILABLE = True
except (ImportError, OSError):
    MAGIC_AVAILABLE = False
    magic = None

try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    DocxDocument = None

try:
    import markdown
    MARKDOWN_AVAILABLE = True
except ImportError:
    MARKDOWN_AVAILABLE = False
    markdown = None

# --- NEW: Add PDF support ---
try:
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False
    fitz = None

# --- NEW: Add ODT support ---
try:
    from odt import ODTParser
    ODT_AVAILABLE = True
except ImportError:
    ODT_AVAILABLE = False
    ODTParser = None


logger = logging.getLogger(__name__)

class DocumentParserError(Exception):
    """Custom exception for document parsing errors."""
    pass

class DocumentParser:
    """A robust multi-format document parser."""
    def __init__(self):
        self.supported_formats = {
            '.txt': 'Plain Text',
            '.md': 'Markdown',
            '.docx': 'Microsoft Word Document',
            '.pdf': 'Portable Document Format',  # <-- NEW
            '.odt': 'OpenDocument Text'          # <-- NEW
        }
        logger.info("DocumentParser initialized.")
        if not DOCX_AVAILABLE:
            logger.warning(".docx parsing disabled. 'python-docx' not found.")
        if not MARKDOWN_AVAILABLE:
            logger.warning(".md parsing disabled. 'markdown' not found.")
        if not PDF_AVAILABLE:
            logger.warning(".pdf parsing disabled. 'PyMuPDF' not found.")
        if not ODT_AVAILABLE:
            logger.warning(".odt parsing disabled. 'odtpy' not found.")


    def detect_file_type(self, file_path: str) -> str:
        if not os.path.exists(file_path):
            raise DocumentParserError(f"File not found: {file_path}")
        
        # Use libmagic first if available, it's more reliable
        if MAGIC_AVAILABLE:
            try:
                mime = magic.from_file(file_path, mime=True)
                if mime == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    return '.docx'
                elif mime == 'application/pdf':
                    return '.pdf'
                elif mime == 'application/vnd.oasis.opendocument.text':
                    return '.odt'
                elif mime == 'text/plain' or mime == 'text/markdown':
                    # For text, still prefer extension for .md
                    _, ext = os.path.splitext(file_path.lower())
                    if ext == '.md':
                        return '.md'
                    return '.txt'
            except Exception:
                pass
        
        # Fallback to extension
        _, ext = os.path.splitext(file_path.lower())
        return ext

    def detect_encoding(self, file_path: str) -> str:
        try:
            with open(file_path, 'rb') as f:
                raw_data = f.read()
                result = chardet.detect(raw_data)
                return result['encoding'] or 'utf-8'
        except Exception as e:
            raise DocumentParserError(f"Error detecting encoding: {str(e)}")

    def parse_docx(self, file_path: str) -> Dict[str, Any]:
        if not DOCX_AVAILABLE:
            raise DocumentParserError("'python-docx' library not available. Cannot parse .docx")
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            core_properties = doc.core_properties
            metadata = {
                'title': core_properties.title,
                'author': core_properties.author,
                'created': core_properties.created,
                'modified': core_properties.modified,
                'paragraph_count': len(doc.paragraphs)
            }
            return {'content': '\n'.join(paragraphs), 'metadata': metadata, 'format': 'docx'}
        except Exception as e:
            raise DocumentParserError(f"Error parsing .docx file: {str(e)}")

    def parse_markdown(self, file_path: str) -> Dict[str, Any]:
        if not MARKDOWN_AVAILABLE:
            raise DocumentParserError("'markdown' library not available. Cannot parse .md")
        try:
            encoding = self.detect_encoding(file_path)
            with codecs.open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            # Simple strip of markdown syntax. Can be improved.
            text_content = re.sub(r'[#*_`]', '', content)
            text_lines = [line.strip() for line in text_content.split('\n') if line.strip()]
            return {
                'content': '\n'.join(text_lines),
                'metadata': {'line_count': len(lines), 'encoding': encoding},
                'format': 'markdown'
            }
        except Exception as e:
            raise DocumentParserError(f"Error parsing .md file: {str(e)}")

    def parse_txt(self, file_path: str) -> Dict[str, Any]:
        try:
            encoding = self.detect_encoding(file_path)
            with codecs.open(file_path, 'r', encoding=encoding) as f:
                content = f.read()
            return {
                'content': content,
                'metadata': {'line_count': len(content.split('\n')), 'encoding': encoding},
                'format': 'txt'
            }
        except Exception as e:
            raise DocumentParserError(f"Error parsing .txt file: {str(e)}")

    # --- NEW: PDF Parser ---
    def parse_pdf(self, file_path: str) -> Dict[str, Any]:
        if not PDF_AVAILABLE:
            raise DocumentParserError("'PyMuPDF' library not available. Cannot parse .pdf")
        try:
            doc = fitz.open(file_path)
            text_content = ""
            for page in doc:
                text_content += page.get_text() + "\n"
            
            metadata = {
                'title': doc.metadata.get('title'),
                'author': doc.metadata.get('author'),
                'page_count': doc.page_count,
                'format': 'pdf'
            }
            doc.close()
            return {'content': text_content, 'metadata': metadata, 'format': 'pdf'}
        except Exception as e:
            raise DocumentParserError(f"Error parsing .pdf file: {str(e)}")

    # --- NEW: ODT Parser ---
    def parse_odt(self, file_path: str) -> Dict[str, Any]:
        if not ODT_AVAILABLE:
            raise DocumentParserError("'odtpy' library not available. Cannot parse .odt")
        try:
            parser = ODTParser()
            content = parser.toString(file_path)
            return {
                'content': content,
                'metadata': {'format': 'odt'},
                'format': 'odt'
            }
        except Exception as e:
            raise DocumentParserError(f"Error parsing .odt file: {str(e)}")


    def parse(self, file_path: str) -> Dict[str, Any]:
        file_type = self.detect_file_type(file_path)
        
        if file_type == '.docx':
            return self.parse_docx(file_path)
        elif file_type == '.md':
            return self.parse_markdown(file_path)
        elif file_type == '.txt':
            return self.parse_txt(file_path)
        elif file_type == '.pdf':  # <-- NEW
            return self.parse_pdf(file_path)
        elif file_type == '.odt':  # <-- NEW
            return self.parse_odt(file_path)
        else:
            try:
                # Fallback: attempt to parse as plain text
                logger.warning(f"Unsupported file type '{file_type}', attempting to parse as .txt")
                return self.parse_txt(file_path)
            except Exception as e:
                raise DocumentParserError(f"Unsupported and un-parseable file format: {file_type} ({e})")

class WebSanitizer:
    """Fetches and sanitizes web content for safe ingestion."""
    
    @staticmethod
    def fetch_and_sanitize(url: str) -> Dict[str, str]:
        """
        Fetches a URL and returns the sanitized plain text.
        
        Args:
            url: The URL to fetch and sanitize
            
        Returns:
            Dictionary with 'content', 'title', and 'source_url'
            
        Raises:
            DocumentParserError: If fetching or sanitizing fails
        """
        try:
            # 1. Fetch with a user-agent to avoid blocking
            headers = {'User-Agent': 'CaudexPro-Agent/1.0'}
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()
            
            # 2. Parse HTML
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # 3. Aggressive Cleaning (The Safety Step)
            # Remove all script, style, meta, iframe, and link tags
            for script in soup(["script", "style", "meta", "iframe", "link", "noscript"]):
                script.decompose()
            
            # 4. Extract Text
            text = soup.get_text(separator='\n')
            
            # 5. Clean Whitespace
            # Collapse multiple newlines into two max (paragraph breaks)
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            clean_text = '\n'.join(chunk for chunk in chunks if chunk)
            
            title = soup.title.string if soup.title else url
            
            logger.info(f"Successfully sanitized content from {url}")
            
            return {
                "content": clean_text,
                "title": title,
                "source_url": url
            }
            
        except requests.exceptions.RequestException as e:
            logger.error(f"Web fetch failed for {url}: {e}")
            raise DocumentParserError(f"Failed to fetch URL: {e}")
        except Exception as e:
            logger.error(f"Web sanitization failed for {url}: {e}")
            raise DocumentParserError(f"Failed to fetch or sanitize URL: {e}")


@dataclass
class TextChunk:
    """Represents a chunk of text with metadata"""
    text: str
    source_file: str
    chapter: Optional[str] = None
    paragraph_start: Optional[int] = None
    paragraph_end: Optional[int] = None
    character_start: Optional[int] = None
    character_end: Optional[int] = None
    entities: List[str] = None

    def __post_init__(self):
        if self.entities is None:
            self.entities = []

class TextChunker:
    """Splits text into chunks while preserving source attribution and metadata"""
    def __init__(self, chunk_size: int = 500, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        logger.info(f"TextChunker initialized with chunk_size={chunk_size}, overlap={chunk_overlap}")

    def chunk_text(self, text: str, source_file: str, chapter: str = None) -> List[TextChunk]:
        try:
            paragraphs = self._split_into_paragraphs(text)
            chunks = self._create_chunks_from_paragraphs(paragraphs, source_file, chapter)
            logger.info(f"Created {len(chunks)} chunks from {source_file}")
            return chunks
        except Exception as e:
            logger.error(f"Failed to chunk text from {source_file}: {e}")
            raise

    def _split_into_paragraphs(self, text: str) -> List[Tuple[str, int, int]]:
        paragraphs = []
        pos = 0
        # Split by one or more newline characters
        for paragraph_text in re.split(r'\n+', text): # Updated regex for flexibility
            if not paragraph_text.strip():
                pos += len(paragraph_text) + 1 # Account for at least one newline
                continue
            
            # Find the actual start position of non-whitespace text
            start_offset = len(paragraph_text) - len(paragraph_text.lstrip())
            start_pos = pos + start_offset
            end_pos = pos + len(paragraph_text)
            
            paragraphs.append((paragraph_text.strip(), start_pos, end_pos))
            pos = end_pos + 1 # Move position to after the current paragraph and its newline
        return paragraphs

    def _create_chunks_from_paragraphs(self, paragraphs: List[Tuple[str, int, int]], 
                                     source_file: str, chapter: str = None) -> List[TextChunk]:
        chunks = []
        current_chunk_text = ""
        current_chunk_start = None
        current_chunk_end = None
        paragraph_start_index = 0
        
        for i, (paragraph_text, para_start, para_end) in enumerate(paragraphs):
            if not paragraph_text:
                continue

            if not current_chunk_text:
                current_chunk_start = para_start
                paragraph_start_index = i
            
            # Check if adding the new paragraph (plus a newline separator) exceeds chunk size
            if len(current_chunk_text) + len(paragraph_text) + 2 > self.chunk_size and current_chunk_text:
                # Save the current chunk
                chunks.append(TextChunk(
                    text=current_chunk_text.strip(),
                    source_file=source_file,
                    chapter=chapter,
                    paragraph_start=paragraph_start_index,
                    paragraph_end=i - 1,
                    character_start=current_chunk_start,
                    character_end=current_chunk_end
                ))
                
                # --- Start new chunk with overlap ---
                # Find a good overlap point (approx self.chunk_overlap)
                overlap_point = max(0, len(current_chunk_text) - self.chunk_overlap)
                # Find the nearest space to not cut words
                overlap_point = current_chunk_text.rfind(' ', 0, overlap_point) + 1
                
                # Get the text for overlap
                overlap_text = current_chunk_text[overlap_point:]
                
                # Find the new paragraph start index for the overlap
                # This is tricky and an approximation, but better than nothing
                overlap_para_start_index = paragraph_start_index
                temp_len = 0
                for j in range(paragraph_start_index, i):
                    temp_len += len(paragraphs[j][0]) + 2 # Add 2 for newlines
                    if temp_len > overlap_point:
                        overlap_para_start_index = j
                        break
                
                # Set new chunk state
                current_chunk_text = overlap_text + "\n\n" + paragraph_text
                current_chunk_start = para_start - len(overlap_text) # Approximate start
                paragraph_start_index = overlap_para_start_index

            else:
                # Just add the paragraph to the current chunk
                current_chunk_text += ("\n\n" if current_chunk_text else "") + paragraph_text
            
            current_chunk_end = para_end
        
        # Add the last remaining chunk
        if current_chunk_text:
            chunks.append(TextChunk(
                text=current_chunk_text.strip(),
                source_file=source_file,
                chapter=chapter,
                paragraph_start=paragraph_start_index,
                paragraph_end=len(paragraphs) - 1,
                character_start=current_chunk_start,
                character_end=current_chunk_end
            ))
        return chunks

    def chunk_with_entity_detection(self, text: str, source_file: str, 
                                  entities: List[Dict[str, Any]], chapter: str = None) -> List[TextChunk]:
        """
        Chunk text and associate entities with chunks based on text positions
        Args:
            text (str): Text to chunk
            source_file (str): Source file identifier
            entities (List[Dict]): List of entity dictionaries with position info
            chapter (str, optional): Chapter/section name
        Returns:
            List[TextChunk]: List of text chunks with associated entities
        """
        # First create basic chunks
        chunks = self.chunk_text(text, source_file, chapter)
        
        # Guard against empty chunks list
        if not chunks:
            return []
            
        # Associate entities with chunks based on position
        for chunk in chunks:
            chunk_entities = []
            for entity in entities:
                # Check if entity position overlaps with chunk position
                if self._positions_overlap(
                    chunk.character_start, chunk.character_end,
                    entity.get('character_start'), entity.get('character_end')
                ):
                    chunk_entities.append(entity['name'])
            chunk.entities = chunk_entities
        return chunks

    def _positions_overlap(self, start1: int, end1: int, start2: int, end2: int) -> bool:
        """
        Check if two character position ranges overlap
        Args:
            start1 (int): Start of first range
            end1 (int): End of first range
            start2 (int): Start of second range
            end2 (int): End of second range
        Returns:
            bool: True if ranges overlap
        """
        # Handle NoneType inputs
        if None in [start1, end1, start2, end2]:
            return False
        return max(start1, start2) <= min(end1, end2)